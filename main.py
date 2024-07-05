import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
import os
from dotenv import load_dotenv


load_dotenv()
openai_api_key = os.getenv('OPENAI_API_KEY')

client = OpenAI(api_key =openai_api_key )


def create_form():
    st.title('Asylum Story Generator from User Input')

    st.header("Part 1: Personal Details")
    full_name = st.text_input("Full Name")
    dob = st.date_input("Date of Birth")
    place_of_birth = st.text_input("Place of Birth")
    current_address = st.text_input("Current Address in Canada")
    phone = st.text_input("Phone")
    email = st.text_input("Email")
    st.header("Part 2: Family Information")
    marital_status = st.selectbox("Marital Status", ["Single", "Married", "Divorced", "Widowed"])
    spouse_name = st.text_input("Spouse’s Name (if applicable)")
    children = st.text_area("Children (Name(s) and Age(s))")

    st.header("Part 3: Education and Occupation")
    education = st.selectbox("Highest Level of Education", ["No formal education", "Primary school", "Secondary school", "College/University", "Other"])
    education_other = st.text_input("If 'Other', please specify") if education == "Other" else ""
    occupation_home_country = st.text_input("Occupation in Home Country")
    occupation_canada = st.text_input("Occupation in Canada (if applicable)")

    st.header("Part 4: Details of Persecution")

    property_dispute = st.text_area("Background of the Property Dispute (Description of property, Cause of dispute)")
    threats = st.text_area("Threats and Incidents (Describe each threat/incident with date, nature, individuals involved)")

    st.header("Part 5: Involvement of Authorities")
    police_reports = st.radio("Were any police reports filed?", ["Yes", "No"])
    police_report_details = st.text_area("If yes, provide details and attach copies") if police_reports == "Yes" else ""
    government_interactions = st.text_area("Describe any interactions with government officials")

    st.header("Part 6: Political and Social Context")
    political_pressure = st.text_area("Describe any political connections or threats")
    social_impact = st.text_area("Describe the impact on your social life and community standing")

    st.header("Part 7: Impact on Life and Decision to Leave")
    personal_impact = st.text_area("Describe the physical, emotional, and psychological toll")
    decision_to_flee = st.text_area("Explain the final events that led to the decision to leave")

    st.header("Part 8: Current Situation and Future Fears")
    current_life = st.text_area("Describe your current living situation and support systems")
    fear_of_returning = st.text_area("Articulate the specific fears associated with returning to your home country")

    st.header("Part 9: Story Elements")
    tags = st.multiselect("Tags", ["Property Dispute", "Political Threat", "Religious Persecution", "Ethnic Violence", "Other"])
    tags_other = st.text_input("If 'Other', please specify") if "Other" in tags else ""
    movies = st.multiselect("Choose Movies for Inspiration (optional)", ["The Kite Runner", "Hotel Rwanda", "Argo", "Other"])
    movies_other = st.text_input("If 'Other', please specify") if "Other" in movies else ""
    story_theme = st.selectbox("Choose the Story Theme", ["Romantic", "Religious Dispute", "Caste/Inter-caste Marriage", "Political Threat", "Other"])
    story_theme_other = st.text_input("If 'Other', please specify") if story_theme == "Other" else ""
    num_incidents = st.number_input("Number of Incidents to Include", min_value=1, step=1)
    length_of_story = st.selectbox("Length of Story", ["Short (1-2 pages)", "Medium (3-5 pages)", "Long (6+ pages)"])
    num_paragraphs = st.number_input("Number of Paragraphs", min_value=1, step=1)

    st.header("Part 10: Conclusion")
    appeal_for_protection = st.text_area("Summarize why you believe you meet the criteria for refugee protection")
    additional_info = st.text_area("Additional Information (optional)")

      

    if st.button("Generate Story"):
            if not all([full_name,place_of_birth,property_dispute]):
              st.warning("Please fill in Name, Place of Birth and One Dispute to generate the story.")
              return 
            input_data = {
                'full_name': full_name,
                'dob': dob,
                'place_of_birth': place_of_birth,
                'current_address': current_address,
                'phone': phone,
                'email': email,
                'marital_status': marital_status,
                'spouse_name': spouse_name,
                'children': children,
                'education': education,
                'education_other': education_other,
                'occupation_home_country': occupation_home_country,
                'occupation_canada': occupation_canada,
                'property_dispute': property_dispute,
                'threats': threats,
                'police_reports': police_reports,
                'police_report_details': police_report_details,
                'government_interactions': government_interactions,
                'political_pressure': political_pressure,
                'social_impact': social_impact,
                'personal_impact': personal_impact,
                'decision_to_flee': decision_to_flee,
                'current_life': current_life,
                'fear_of_returning': fear_of_returning,
                'tags': tags,
                'tags_other': tags_other,
                'movies': movies,
                'movies_other': movies_other,
                'story_theme': story_theme,
                'story_theme_other': story_theme_other,
                'num_incidents': num_incidents,
                'length_of_story': length_of_story,
                'num_paragraphs': num_paragraphs,
                'appeal_for_protection': appeal_for_protection,
                'additional_info': additional_info,
            }

            story = generate_story(input_data)
            st.subheader('Generated Story')
            st.write(story)
      

            if story:
                doc = Document()
                doc.add_heading('Asylum Story', 0)

                for line in story.split('\n'):
                    doc.add_paragraph(line)

                bio = BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.download_button(
                    label="Download Story as Word Document",
                    data=bio,
                    file_name="asylum_story.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def generate_story(details):
    prompt = [
        {"role": "assistant", "content": "You are a Journalist. Your work is to create awareness about people seeking refuge in different countries and their reasons behind this. Like a good Journalist, use the user information to create an impactful story which would help the refugee to seek better opportunities."},
        {"role": "user", "content":  f"A refugee named {details['full_name']}, aged {details['dob']} from {details['place_of_birth']}, currently residing at {details['current_address']}. "f"They are {details['marital_status']} with {details['children']} children. Their education level is {details['education']} "f"and occupation in their home country is {details['occupation_home_country']}. They faced a property dispute related to {details['property_dispute']} "f"and encountered threats like {details['threats']}. They interacted with government officials regarding {details['government_interactions']} "f"and felt {details['political_pressure']} political pressure. Their current life situation involves {details['current_life']} and they fear returning to {details['fear_of_returning']}.Their story theme is {details['story_theme']}." f"{details['additional_info'] if details['additional_info'] else ''}"
   
        }
    ]


    response = client.chat.completions.create(model="gpt-3.5-turbo",  # Use the desired model
    messages=prompt,
    max_tokens=3400,
    n=4,
    stop=None,
    temperature=0.7)

    story=response.choices[0].message.content
    return story

if __name__ == "__main__":
    create_form()
    
